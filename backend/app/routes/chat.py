import os
import json
import re
import unicodedata
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel

from ..config import settings
from ..services.renderer import renderer
from ..services.auth_wallet import (
    AuthenticatedUser,
    InsufficientBalanceError,
    credit_balance,
    debit_balance,
    list_user_chat_history,
    record_chat_history,
    require_authenticated_user,
)
from ..services.billing import module_cost_chat_cop, module_cost_chat_image_cop

router = APIRouter(prefix="/chat", tags=["chat"])
CHAT_EXPORTS_DIR = Path(settings.data_dir) / "chat_exports"
MAX_CHAT_ATTACHMENTS = 5
MAX_CHAT_ATTACHMENT_BYTES = 10 * 1024 * 1024
SAFE_ATTACHMENT_RE = re.compile(r"[^a-zA-Z0-9._-]+")
TEXT_ATTACHMENT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".log", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".html", ".css", ".xml", ".yml", ".yaml", ".ini", ".toml", ".sql",
}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}
MAX_ATTACHMENT_TEXT_CHARS = 4000

IA_IMP_DEFINITION_RESPONSE = "IA-IMP es una creación 100% de IMPORMADERAS Y que es IMPORMADERAS"

IMPORMADERAS_DEFINITION_RESPONSE = (
    "Somos una empresa caleña con más de 20 años de trayectoria en el sector maderero y mobiliario, especializada en el diseño, arquitectura y distribución de materiales de alta calidad para la fabricación y transformación de espacios. Nuestro compromiso con el desarrollo regional y la innovación nos ha permitido consolidarnos como un aliado estratégico para profesionales, empresas y proyectos de interiorismo en todo el suroccidente colombiano.\n\n"
    "Contamos con un portafolio integral que incluye laminados RH y estándar, MDF, Triplex y una amplia variedad de tableros decorativos, así como herrajes, accesorios, soluciones en ferretería y elementos para diseño de interiores. Cada producto es seleccionado bajo criterios de durabilidad, estética y funcionalidad, garantizando resultados óptimos en cada proyecto.\n\n"
    "Nuestra experiencia, servicio cercano y conocimiento del mercado nos permiten ofrecer asesoría especializada, acompañamiento técnico y soluciones personalizadas para arquitectos, diseñadores, constructores, carpinterías y clientes finales que buscan elevar la calidad de sus espacios.\n\n"
    "En Impormaderas trabajamos con pasión, responsabilidad y visión a futuro, promoviendo el buen diseño, el trabajo en equipo y el crecimiento de nuestra región. Aquí, cada proyecto importa y cada detalle cuenta.**Somos más que materiales: somos aliados en la construcción de ideas.**"
)

COUNTRY_CITY_RESPONSE = "Colombia - Exactamente en Cali - Palmira - Dosquebradas y Jamundí"

HEADQUARTERS_RESPONSE = (
    "Dando clic a este link podras consultar todas nuestras sedes, incluso tendras la ruta con Waze: "
    "https://tiendaonline.impormaderasltda.com/sedes/"
)

CONTACT_LINES_RESPONSE = (
    "Si necesitas un numero de atención mas agil, comunicate a la linea de atención digital:\n"
    "WhatsApp: 3137399382\n"
    "Es más te doy el enlace directo, solo dale clic y te llevamos al WhatsApp https://wa.link/n32n25"
)

WEBSITE_RESPONSE = "La pagina web oficial de IMPORMADERAS es: https://impormaderasltda.com/"

RECENT_ATTACHMENTS_BY_USER: dict[int, list[dict]] = {}


class ChatRequest(BaseModel):
    message: str
    context: str = ""
    conversation_id: str = ""


def _safe_attachment_name(file_name: str) -> str:
    original = Path(str(file_name or "archivo")).name
    stem = SAFE_ATTACHMENT_RE.sub("_", Path(original).stem)[:60] or "archivo"
    suffix = Path(original).suffix.lower()
    if not suffix:
        suffix = ".bin"
    if len(suffix) > 10:
        suffix = ".bin"
    return f"upload-{uuid4()}-{stem}{suffix}"


async def _save_chat_attachments(files: list[object]) -> list[dict]:
    if not files:
        return []

    if len(files) > MAX_CHAT_ATTACHMENTS:
        raise HTTPException(status_code=400, detail=f"Puedes adjuntar hasta {MAX_CHAT_ATTACHMENTS} archivos por mensaje")

    CHAT_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    saved: list[dict] = []

    for item in files:
        if not hasattr(item, "read"):
            continue
        file_name = str(getattr(item, "filename", "") or "")
        if not file_name:
            continue

        payload = await item.read()
        if len(payload) > MAX_CHAT_ATTACHMENT_BYTES:
            raise HTTPException(status_code=400, detail="Cada archivo debe pesar maximo 10 MB")

        text_content = _extract_attachment_text(file_name, payload)

        stored_name = _safe_attachment_name(file_name)
        target = _safe_export_path(stored_name)
        target.write_bytes(payload)

        saved.append(
            {
                "original_name": file_name,
                "stored_name": stored_name,
                "size": len(payload),
                "content_type": str(getattr(item, "content_type", "") or "application/octet-stream"),
                "url": f"/chat/files/{stored_name}",
                "text_content": text_content,
            }
        )

    return saved


def _extract_attachment_text(file_name: str, payload: bytes) -> str:
    suffix = Path(str(file_name or "")).suffix.lower()

    def _truncate(text: str) -> str:
        return (text or "")[:MAX_ATTACHMENT_TEXT_CHARS]

    if suffix in TEXT_ATTACHMENT_EXTENSIONS:
        try:
            return _truncate(payload.decode("utf-8", errors="replace"))
        except Exception:
            return ""

    if suffix in SPREADSHEET_EXTENSIONS:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(filename=BytesIO(payload), data_only=True, read_only=True)
            chunks: list[str] = []
            for sheet in wb.worksheets[:3]:
                chunks.append(f"Hoja: {sheet.title}")
                row_count = 0
                for row in sheet.iter_rows(min_row=1, max_row=80, values_only=True):
                    cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
                    if not cells:
                        continue
                    chunks.append(" | ".join(cells))
                    row_count += 1
                    if row_count >= 40:
                        break
            return _truncate("\n".join(chunks))
        except Exception:
            return ""

    if suffix in PDF_EXTENSIONS:
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(payload))
            pages_text: list[str] = []
            for idx, page in enumerate(reader.pages[:8], start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    pages_text.append(f"Pagina {idx}: {page_text}")
            if pages_text:
                return _truncate("\n\n".join(pages_text))
            return "PDF adjunto detectado, pero no se pudo extraer texto legible automaticamente."
        except Exception:
            return "PDF adjunto detectado. Si deseas analisis de contenido, instala pypdf en backend."

    if suffix in IMAGE_EXTENSIONS:
        try:
            from PIL import Image

            with Image.open(BytesIO(payload)) as img:
                base = f"Imagen adjunta: formato={img.format or 'desconocido'}, tamaño={img.width}x{img.height}, modo={img.mode}."
                try:
                    import pytesseract

                    ocr_text = (pytesseract.image_to_string(img) or "").strip()
                    if ocr_text:
                        return _truncate(f"{base}\nTexto detectado por OCR:\n{ocr_text}")
                    return _truncate(f"{base}\nNo se detecto texto legible por OCR.")
                except Exception:
                    return _truncate(f"{base}\nOCR no disponible en servidor (instalar pytesseract + motor Tesseract).")
        except Exception:
            return "Imagen adjunta detectada, pero no se pudo leer el contenido de la imagen."

    if suffix == ".docx":
        try:
            from docx import Document

            doc = Document(BytesIO(payload))
            text = "\n".join((p.text or "") for p in doc.paragraphs)
            return _truncate(text)
        except Exception:
            return ""

    return f"Adjunto {file_name} detectado (tipo {suffix or 'desconocido'})."


def _build_attachment_context(attachments: list[dict]) -> str:
    if not attachments:
        return ""
    lines = ["Adjuntos del usuario:"]
    for index, item in enumerate(attachments, start=1):
        name = str(item.get("original_name") or "archivo")
        mime = str(item.get("content_type") or "application/octet-stream")
        size = int(item.get("size") or 0)
        url = str(item.get("url") or "")
        lines.append(f"{index}. {name} | tipo={mime} | bytes={size} | enlace={url}")
        content_preview = str(item.get("text_content") or "").strip()
        if content_preview:
            lines.append(f"Contenido detectado en {name}:\n{content_preview}")
    return "\n".join(lines)


def _looks_like_excel_request(text: str) -> bool:
    normalized = text.lower()
    keywords = ["excel", "xlsx", "cotizacion", "cotización", "plantilla", "descargable", "hoja de calculo"]
    return sum(1 for k in keywords if k in normalized) >= 2


def _looks_like_image_request(text: str) -> bool:
    normalized = text.lower()
    keywords = [
        "logo",
        "imagen",
        "diseña",
        "disena",
        "render",
        "mockup",
        "banner",
        "afiche",
        "flyer",
        "poster",
        "portada",
        "branding",
        "visual",
    ]
    return any(k in normalized for k in keywords)


def _looks_like_word_request(text: str) -> bool:
    normalized = text.lower()
    keywords = [
        "word",
        "docx",
        "documento",
        "documentos",
        "carta",
        "oficio",
        "informe",
        "memo",
        "memorando",
        "solicitud",
        "acta",
        "contrato",
        "certificado",
    ]
    triggers = ["crear", "genera", "generar", "haz", "prepara", "redacta", "descargable"]
    return any(k in normalized for k in keywords) and any(t in normalized for t in triggers)


def _normalize_intent_text(text: str) -> str:
    lowered = (text or "").strip().lower()
    normalized = unicodedata.normalize("NFKD", lowered)
    ascii_only = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9\s]", " ", ascii_only)


def _get_ia_imp_default_answer(message: str) -> str | None:
    normalized = _normalize_intent_text(message)
    compact = " ".join(normalized.split())

    if "que es ia imp" in compact:
        return IA_IMP_DEFINITION_RESPONSE

    if "que es impormaderas" in compact or "quien es impormaderas" in compact:
        return IMPORMADERAS_DEFINITION_RESPONSE

    if "en que pais estan" in compact or "en que ciudad estan" in compact:
        return COUNTRY_CITY_RESPONSE

    asks_for_locations = any(
        phrase in compact
        for phrase in (
            "donde estan sus sedes",
            "donde estan las sedes",
            "cuales son las sedes",
            "cuales son sus sedes",
            "sedes de impormaderas",
            "ubicacion de las sedes",
            "ubicaciones de impormaderas",
        )
    )
    mentions_sedes = "sede" in compact or "sedes" in compact
    mentions_impormaderas = "impormaderas" in compact
    if asks_for_locations or (mentions_sedes and ("donde" in compact or "cual" in compact or mentions_impormaderas)):
        return HEADQUARTERS_RESPONSE

    asks_for_website = any(
        phrase in compact
        for phrase in (
            "pagina web de impormaderas",
            "sitio web de impormaderas",
            "web de impormaderas",
            "pagina de impormaderas",
            "cual es la pagina web",
            "cual es el sitio web",
            "dominio de impormaderas",
        )
    )
    if asks_for_website:
        return WEBSITE_RESPONSE

    if "lineas de atencion" in compact or "linea de atencion" in compact:
        return CONTACT_LINES_RESPONSE

    return None


def _remember_user_attachments(user_id: int, attachments: list[dict]) -> None:
    if not attachments:
        return
    RECENT_ATTACHMENTS_BY_USER[int(user_id)] = [dict(item) for item in attachments]


def _get_user_recent_attachments(user_id: int) -> list[dict]:
    items = RECENT_ATTACHMENTS_BY_USER.get(int(user_id), [])
    return [dict(item) for item in items]


def _is_image_attachment(item: dict) -> bool:
    content_type = str(item.get("content_type") or "").lower()
    if content_type.startswith("image/"):
        return True
    suffix = Path(str(item.get("original_name") or "")).suffix.lower()
    return suffix in IMAGE_EXTENSIONS


def _first_image_attachment_path(attachments: list[dict]) -> Path | None:
    for item in attachments:
        if not isinstance(item, dict) or not _is_image_attachment(item):
            continue
        stored_name = str(item.get("stored_name") or "").strip()
        if not stored_name:
            continue
        candidate = _safe_export_path(stored_name)
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def _looks_like_image_transform_request(text: str) -> bool:
    normalized = _normalize_intent_text(text)
    compact = " ".join(normalized.split())
    if not compact:
        return False

    action_terms = [
        "genera",
        "generar",
        "crea",
        "crear",
        "haz",
        "usa",
        "utiliza",
        "aplica",
        "coloca",
        "pon",
        "integra",
        "incorpora",
        "transforma",
        "transformar",
        "mejora",
        "editar",
        "edita",
        "reimagina",
        "redisena",
        "disena",
    ]
    reference_terms = [
        "con esta imagen",
        "con esa imagen",
        "este logo",
        "ese logo",
        "con el logo adjunto",
        "con logo adjunto",
        "logo adjunto",
        "archivo adjunto",
        "con este archivo",
        "con ese archivo",
        "archivo que adjunte",
        "archivo que subi",
        "logo que te comparto",
        "te comparto",
        "adjunto",
        "adjuntos",
        "logo que adjunte",
        "logo que subi",
        "esta imagen",
        "esa imagen",
        "esta foto",
        "esa foto",
        "imagen adjunta",
        "foto adjunta",
        "imagen que adjunte",
        "foto que adjunte",
        "imagen que subi",
        "foto que subi",
        "la imagen",
        "la foto",
        "basado en",
        "a partir de",
    ]

    has_action = any(term in compact for term in action_terms)
    has_reference_hint = any(term in compact for term in reference_terms)
    return has_action and has_reference_hint


def _save_chat_model_output(output: object, output_image_path: str) -> None:
    target = Path(output_image_path)
    target.parent.mkdir(parents=True, exist_ok=True)

    if hasattr(output, "read"):
        target.write_bytes(output.read())
        return

    if hasattr(output, "url"):
        remote_url = str(output.url)
    else:
        remote_url = str(output)

    if remote_url.startswith("http://") or remote_url.startswith("https://"):
        import urllib.request

        with urllib.request.urlopen(remote_url, timeout=120) as response:
            target.write_bytes(response.read())
        return

    raise RuntimeError("No se pudo interpretar la salida del modelo de imagen")


def _generate_from_reference_image(reference_image_path: str, prompt_text: str, output_image_path: str) -> str:
    if not settings.replicate_api_token:
        raise RuntimeError("Falta REPLICATE_API_TOKEN en backend/.env")

    os.environ["REPLICATE_API_TOKEN"] = settings.replicate_api_token
    import replicate

    guidance_prompt = (
        f"{prompt_text.strip()}\n"
        "Usa la imagen adjunta como referencia principal y conserva de forma fiel su logo o marca. "
        "Si el usuario pide un mockup o escena (ej: vaso en taller), integra el logo de la referencia sobre el objeto solicitado."
    ).strip()

    ref_path = Path(reference_image_path)
    attempts = [
        {
            "prompt": guidance_prompt,
            "image_input": [ref_path],
            "aspect_ratio": "1:1",
            "output_format": "png",
            "prompt_upsampling": True,
            "safety_tolerance": 2,
        },
        {
            "prompt": guidance_prompt,
            "input_image": [ref_path],
            "aspect_ratio": "1:1",
            "output_format": "png",
            "prompt_upsampling": True,
            "safety_tolerance": 2,
        },
    ]

    last_error: Exception | None = None
    output: object | None = None

    for payload in attempts:
        try:
            output = replicate.run("google/nano-banana", input=payload)
            break
        except Exception as exc:
            last_error = exc

    if output is None:
        raise RuntimeError(f"No se pudo generar con referencia adjunta: {last_error}")

    first_output = output[0] if isinstance(output, list) and output else output
    _save_chat_model_output(first_output, output_image_path)
    return "google/nano-banana"


def _safe_export_path(file_name: str) -> Path:
    candidate = (CHAT_EXPORTS_DIR / file_name).resolve()
    base = CHAT_EXPORTS_DIR.resolve()
    try:
        candidate.relative_to(base)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Archivo invalido") from exc
    return candidate


def _build_excel_template(file_path: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.worksheet.datavalidation import DataValidation
    except Exception as exc:
        raise RuntimeError("Falta dependencia openpyxl para generar Excel") from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "Cotizacion"

    ws["A1"] = "Plantilla de Cotizacion - Carpinteria"
    ws.merge_cells("A1:H1")
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    headers = [
        "Item",
        "Descripcion",
        "Cantidad",
        "Unidad",
        "Precio Unitario",
        "Descuento %",
        "Subtotal",
        "Total",
    ]

    header_fill = PatternFill(start_color="1FC86C", end_color="1FC86C", fill_type="solid")
    header_font = Font(color="000000", bold=True)
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, name in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=col, value=name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 42
    ws.column_dimensions["C"].width = 11
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 16
    ws.column_dimensions["F"].width = 13
    ws.column_dimensions["G"].width = 16
    ws.column_dimensions["H"].width = 16

    unit_validation = DataValidation(type="list", formula1='"Und,ml,m2,m3,global"', allow_blank=True)
    ws.add_data_validation(unit_validation)

    for row in range(4, 54):
        ws.cell(row=row, column=1, value=row - 3)
        ws.cell(row=row, column=7, value=f"=C{row}*E{row}")
        ws.cell(row=row, column=8, value=f"=G{row}*(1-F{row})")

        for col in range(1, 9):
            ws.cell(row=row, column=col).border = border

        ws.cell(row=row, column=3).number_format = "0.00"
        ws.cell(row=row, column=5).number_format = '#,##0.00'
        ws.cell(row=row, column=6).number_format = "0.00%"
        ws.cell(row=row, column=7).number_format = '#,##0.00'
        ws.cell(row=row, column=8).number_format = '#,##0.00'

        unit_validation.add(ws.cell(row=row, column=4))

    ws["G56"] = "Subtotal"
    ws["H56"] = "=SUM(H4:H53)"
    ws["G57"] = "IVA 19%"
    ws["H57"] = "=H56*0.19"
    ws["G58"] = "TOTAL"
    ws["H58"] = "=H56+H57"

    for row in (56, 57, 58):
        ws[f"G{row}"].font = Font(bold=True)
        ws[f"H{row}"].font = Font(bold=True)
        ws[f"G{row}"].border = border
        ws[f"H{row}"].border = border
        ws[f"H{row}"].number_format = '#,##0.00'

    ws.freeze_panes = "A4"

    CHAT_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(file_path)


def _generate_logo_image(prompt_text: str) -> Path:
    CHAT_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    file_path = CHAT_EXPORTS_DIR / f"logo-{uuid4()}.png"

    prompt = (
        "Diseña un logo moderno, limpio y profesional para carpinteria. "
        "Fondo blanco, composicion centrada, alta legibilidad, estilo vectorial limpio, "
        "sin mockup ni sombras exageradas. Instruccion del cliente: "
        f"{prompt_text.strip()}"
    )

    renderer.generate_text_to_image(
        output_image_path=str(file_path),
        prompt=prompt,
        negative_prompt="low quality, blurry, watermark, mockup, distorted text",
        steps=4,
        guidance_scale=7.0,
        quality="balanced",
        seed=None,
    )
    return file_path


def _build_word_document(prompt_text: str) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("Falta dependencia python-docx para generar Word") from exc

    CHAT_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    file_path = CHAT_EXPORTS_DIR / f"documento-{uuid4()}.docx"

    prompt = prompt_text.strip()
    prompt_lower = prompt.lower()

    document = Document()
    title = "Documento generado por Pachy IA"
    if "carta" in prompt_lower:
        title = "Carta generada por Pachy IA"

    document.add_heading(title, level=1)
    document.add_paragraph("Borrador editable creado a partir de la solicitud del usuario.")
    document.add_paragraph(f"Solicitud: {prompt}")

    intro = document.add_paragraph()
    intro.add_run("Introduccion: ").bold = True
    intro.add_run(
        "Este documento fue generado como una base profesional para que puedas revisar, editar y enviar. "
        "Si necesitas formato legal, academico o comercial mas estricto, conviene revisar datos y firmas antes de usarlo."
    )

    if "carta" in prompt_lower:
        document.add_paragraph("Destinatario: ______________________________")
        document.add_paragraph("Asunto: ______________________________")
        document.add_paragraph(
            "Cuerpo: redacta aqui el mensaje principal con tono formal, claro y directo, adaptado al objetivo de la carta."
        )
        document.add_paragraph("Despedida: Atentamente,")
        document.add_paragraph("Nombre y cargo: ______________________________")
        document.add_paragraph("Firma: ______________________________")
    else:
        document.add_paragraph("Objetivo del documento: ______________________________")
        document.add_paragraph("Contenido principal: desarrolla aqui los puntos clave solicitados por el usuario.")
        document.add_paragraph("Conclusiones o siguientes pasos: ______________________________")

    footer = document.add_paragraph()
    footer_run = footer.add_run("Pachy IA puede usar este archivo como base para cartas, solicitudes, informes y documentos internos.")
    footer_run.italic = True

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    document.save(file_path)
    return file_path


@router.get("/files/{file_name}")
def get_chat_file(file_name: str):
    path = _safe_export_path(file_name)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="Archivo no encontrado")

    media_type = "application/octet-stream"
    if path.suffix.lower() == ".xlsx":
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif path.suffix.lower() == ".docx":
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
        media_type = f"image/{'jpeg' if path.suffix.lower() in {'.jpg', '.jpeg'} else path.suffix.lower().lstrip('.')}"

    return FileResponse(path=path, media_type=media_type, filename=path.name)


def _extract_text(output: object) -> str:
    if output is None:
        return ""

    if isinstance(output, str):
        return output.strip()

    if isinstance(output, list):
        parts: list[str] = []
        for item in output:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict):
                txt = item.get("text") or item.get("output") or item.get("content")
                if txt:
                    parts.append(str(txt))
            else:
                parts.append(str(item))

        # Some Replicate chat models stream token/chunk arrays.
        # Joining with empty string avoids artificial line breaks between tokens.
        merged = "".join([p for p in parts if p]).strip()
        # Normalize excessive spaces without altering paragraph breaks.
        return "\n".join(" ".join(line.split()) for line in merged.splitlines())

    if isinstance(output, dict):
        # OpenAI-like response shape.
        choices = output.get("choices") if isinstance(output.get("choices"), list) else None
        if choices:
            first = choices[0] if choices else None
            if isinstance(first, dict):
                message = first.get("message")
                if isinstance(message, dict):
                    content = message.get("content")
                    if isinstance(content, str) and content.strip():
                        return content.strip()
                text = first.get("text")
                if isinstance(text, str) and text.strip():
                    return text.strip()

        for key in ("output", "text", "response", "content"):
            if key in output and output[key]:
                return str(output[key]).strip()

    return str(output).strip()


def _serialize_chat_attachments(items: list[dict]) -> list[dict]:
    safe_items: list[dict] = []
    for raw in items or []:
        if not isinstance(raw, dict):
            continue
        safe_items.append(
            {
                "original_name": str(raw.get("original_name") or ""),
                "content_type": str(raw.get("content_type") or ""),
                "size": int(raw.get("size") or 0),
                "url": str(raw.get("url") or ""),
            }
        )
    return safe_items


def _build_chat_response(
    *,
    user: AuthenticatedUser,
    user_message: str,
    answer: str,
    model: str,
    conversation_id: str,
    attachments: list[dict],
) -> dict:
    try:
        record_chat_history(
            user_id=user.user_id,
            conversation_id=conversation_id,
            user_message=user_message,
            assistant_message=answer,
            model=model,
            attachments=_serialize_chat_attachments(attachments),
        )
    except Exception:
        # History persistence should not block the user-facing chat response.
        pass

    return {
        "answer": answer,
        "model": model,
        "attachments": attachments,
    }


@router.get("/history", response_model=dict)
def get_history(user: AuthenticatedUser = Depends(require_authenticated_user)):
    return {"items": list_user_chat_history(user.user_id, limit=200)}


@router.post("/message", response_model=dict)
async def chat_message(request: Request, user: AuthenticatedUser = Depends(require_authenticated_user)):
    content_type = str(request.headers.get("content-type") or "").lower()
    message = ""
    context = ""
    conversation_id = ""
    attachments: list[dict] = []

    if content_type.startswith("multipart/form-data"):
        form = await request.form()
        message = str(form.get("message") or "").strip()
        context = str(form.get("context") or "")
        conversation_id = str(form.get("conversation_id") or "").strip()
        files = list(form.getlist("files"))
        attachments = await _save_chat_attachments(files)
    else:
        payload_raw = await request.json()
        payload = ChatRequest(**payload_raw)
        message = payload.message.strip()
        context = payload.context
        conversation_id = payload.conversation_id.strip()

    if not conversation_id:
        conversation_id = f"conv-{uuid4()}"

    if attachments:
        _remember_user_attachments(user.user_id, attachments)
    effective_attachments = attachments or _get_user_recent_attachments(user.user_id)
    persisted_user_message = message.strip() or "Mensaje con adjuntos"

    if not message and not effective_attachments:
        raise HTTPException(status_code=400, detail="El mensaje no puede estar vacio")

    reference_image = _first_image_attachment_path(effective_attachments)
    wants_chat_image = False
    if reference_image is not None and (
        _looks_like_image_transform_request(message) or _looks_like_image_request(message)
    ):
        wants_chat_image = True
    elif _looks_like_image_request(message):
        wants_chat_image = True

    billed_amount = module_cost_chat_image_cop() if wants_chat_image else module_cost_chat_cop()
    debit_module = "chat_image" if wants_chat_image else "chat"
    debit_note = "Generacion de imagen en Pachy IA" if wants_chat_image else "Consulta Pachy IA"
    try:
        debit_balance(user.user_id, billed_amount, debit_module, debit_note)
    except InsufficientBalanceError as exc:
        if wants_chat_image:
            raise HTTPException(
                status_code=402,
                detail="Saldo insuficiente. Recarga tu cuenta para generar imagenes con Pachy IA.",
            ) from exc
        raise HTTPException(status_code=402, detail="Saldo insuficiente. Recarga tu cuenta para usar Pachy IA.") from exc

    try:
        default_answer = _get_ia_imp_default_answer(message)
        if default_answer:
            return _build_chat_response(
                user=user,
                user_message=persisted_user_message,
                answer=default_answer,
                model="policy:ia-imp-default",
                conversation_id=conversation_id,
                attachments=effective_attachments,
            )

        if reference_image is not None and (
            _looks_like_image_transform_request(message) or _looks_like_image_request(message)
        ):
            file_name = f"chat-image-{uuid4()}.png"
            output_path = _safe_export_path(file_name)
            prompt_text = message or "Genera una version mejorada basada en la imagen adjunta"
            used_model = "tool:img2img-from-attachment"
            try:
                used_model = _generate_from_reference_image(
                    reference_image_path=str(reference_image),
                    prompt_text=prompt_text,
                    output_image_path=str(output_path),
                )
            except Exception:
                renderer.generate(
                    input_image_path=str(reference_image),
                    output_image_path=str(output_path),
                    prompt=prompt_text,
                    negative_prompt="low quality, blurry, watermark, deformed",
                    steps=30,
                    guidance_scale=7.0,
                    quality="balanced",
                    seed=None,
                )
            return _build_chat_response(
                user=user,
                user_message=persisted_user_message,
                answer=(
                    "Listo, usé tu imagen adjunta como referencia para generar el resultado. "
                    f"Descarga aquí: /chat/files/{file_name}"
                ),
                model=used_model,
                conversation_id=conversation_id,
                attachments=effective_attachments,
            )

        if _looks_like_excel_request(message):
            file_name = f"cotizacion-{uuid4()}.xlsx"
            file_path = _safe_export_path(file_name)
            _build_excel_template(file_path)
            return _build_chat_response(
                user=user,
                user_message=persisted_user_message,
                answer=(
                    "Listo, generé tu Excel de cotización con fórmulas y formato listo para reemplazar valores. "
                    f"Descarga aquí: /chat/files/{file_name}"
                ),
                model="tool:excel-template",
                conversation_id=conversation_id,
                attachments=effective_attachments,
            )

        if _looks_like_word_request(message):
            generated = _build_word_document(message)
            return _build_chat_response(
                user=user,
                user_message=persisted_user_message,
                answer=(
                    "Listo, generé tu documento en Word con un borrador editable. "
                    f"Descarga aquí: /chat/files/{generated.name}"
                ),
                model="tool:word-doc",
                conversation_id=conversation_id,
                attachments=effective_attachments,
            )

        if _looks_like_image_request(message):
            if not settings.replicate_api_token:
                raise HTTPException(status_code=500, detail="Falta REPLICATE_API_TOKEN en backend/.env")
            generated = _generate_logo_image(message)
            return _build_chat_response(
                user=user,
                user_message=persisted_user_message,
                answer=(
                    "Listo, generé una propuesta visual descargable. "
                    f"Descarga aquí: /chat/files/{generated.name}"
                ),
                model="tool:image-generator",
                conversation_id=conversation_id,
                attachments=effective_attachments,
            )

        if not settings.replicate_api_token:
            raise HTTPException(status_code=500, detail="Falta REPLICATE_API_TOKEN en backend/.env")

        os.environ["REPLICATE_API_TOKEN"] = settings.replicate_api_token

        import replicate

        system_prefix = (
            "Eres IA-IMP, asistente arquitectonico de IMPORMADERAS. "
            "Responde en espanol claro, tecnico y practico."
        )
        attachment_context = _build_attachment_context(effective_attachments)
        context_block = context
        if attachment_context:
            context_block = f"{context}\n\n{attachment_context}".strip()
        user_prompt = f"Contexto:\n{context_block}\n\nUsuario:\n{message}".strip()
        composed_prompt = f"{system_prefix}\n\n{user_prompt}\n\nAsistente:".strip()

        model_candidates = [
            settings.replicate_chat_model,
            "openai/gpt-4.1-mini",
            "anthropic/claude-3.7-sonnet",
            "meta/meta-llama-3-70b-instruct",
        ]

        # Remove duplicates while preserving order.
        unique_models: list[str] = []
        for model in model_candidates:
            safe_model = str(model or "").strip()
            if safe_model and safe_model not in unique_models:
                unique_models.append(safe_model)

        last_error: Exception | None = None
        output: object | None = None
        used_model: str | None = None

        for model_id in unique_models:
            attempts = [
                {
                    "system_prompt": system_prefix,
                    "prompt": user_prompt,
                    "temperature": 0.5,
                    "max_completion_tokens": 4096,
                },
                {
                    "messages": json.dumps(
                        [
                            {"role": "system", "content": system_prefix},
                            {"role": "user", "content": user_prompt},
                        ]
                    ),
                    "temperature": 0.5,
                    "max_completion_tokens": 4096,
                },
                {"prompt": composed_prompt},
                {"input": composed_prompt},
                {"message": composed_prompt},
            ]

            for inp in attempts:
                try:
                    output = replicate.run(model_id, input=inp)
                    used_model = model_id
                    break
                except Exception as exc:
                    last_error = exc

            if output is not None:
                break

        if output is None:
            raise RuntimeError(f"Replicate chat fallo: {last_error}")

        answer = _extract_text(output)
        if not answer:
            answer = "No pude generar respuesta en este momento."

        return _build_chat_response(
            user=user,
            user_message=persisted_user_message,
            answer=answer,
            model=str(used_model or ""),
            conversation_id=conversation_id,
            attachments=effective_attachments,
        )
    except Exception as exc:
        try:
            refund_module = "chat_image_refund" if wants_chat_image else "chat_refund"
            refund_note = "Reembolso por fallo en imagen Pachy IA" if wants_chat_image else "Reembolso por fallo en Pachy IA"
            credit_balance(user.user_id, billed_amount, refund_module, refund_note)
        except Exception:
            pass
        if isinstance(exc, HTTPException):
            raise exc
        raise HTTPException(status_code=500, detail=str(exc)) from exc
